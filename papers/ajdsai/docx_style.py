"""Post-process the pandoc-generated DOCX so it matches the LaTeX PDF style:

  1. Number figures and tables in document order (Figure 1.., Table 1..),
     reproducing the LaTeX counters, with the label in bold (labelfont=bf).
  2. Apply a booktabs look to every table: thick top rule, thin rule under
     the header row, thick bottom rule, no vertical lines, no inner
     horizontal lines (matches \\toprule/\\midrule/\\bottomrule).
  3. Centre figures and their captions; centre tables.
  4. Prepend the editorial submission note (read from main.tex's
     \\submitnote{} so there is a single source of truth) as the very
     first line of the document, above the title — pandoc cannot expand
     the custom macro, so it is injected here.

Run:  python docx_style.py main.docx
"""
from __future__ import annotations
import re
import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


def extract_submitnote(tex_path: str = "main.tex") -> str | None:
    """Pull the plain text of \\submitnote{...} from the LaTeX source."""
    p = Path(tex_path)
    if not p.exists():
        return None
    text = p.read_text(encoding="utf-8")
    # Match the *usage* (after \begin{document}), not the \newcommand definition.
    for m in re.finditer(r"\\submitnote\{([^}]*)\}", text):
        body = m.group(1)
        if "#1" in body:          # this is the \newcommand definition — skip
            continue
        # Strip leading formatting commands (\small \itshape ...)
        body = re.sub(r"\\[a-zA-Z]+\s*", "", body).strip()
        if body:
            return body
    return None


def prepend_note(doc: Document, text: str) -> None:
    """Insert `text` (small italic) as the first paragraph, above the title."""
    para = doc.add_paragraph()          # appended at end for now
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Move the new paragraph to the very top of the document body.
    body = doc.element.body
    body.insert(0, para._p)


def prepend_bold_label(paragraph: Paragraph, label: str) -> None:
    """Insert `label` (bold) at the very start of the paragraph text."""
    run = paragraph.add_run(label)
    run.bold = True
    r = run._r
    p = paragraph._p
    p.remove(r)
    ppr = p.find(qn("w:pPr"))
    if ppr is not None:
        ppr.addnext(r)
    else:
        p.insert(0, r)


def _border(tag: str, sz: int, val: str = "single"):
    el = OxmlElement(f"w:{tag}")
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), "000000")
    return el


def apply_booktabs(table: Table) -> None:
    """Thick top/bottom rules, thin rule under header, no verticals/inner lines."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    # Clear any pandoc-applied table style so our borders win
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    borders.append(_border("top", 14))           # ~1.75pt thick top rule
    borders.append(_border("bottom", 14))        # thick bottom rule
    borders.append(_border("left", 0, "nil"))
    borders.append(_border("right", 0, "nil"))
    borders.append(_border("insideH", 0, "nil"))  # no inner horizontal lines
    borders.append(_border("insideV", 0, "nil"))  # no vertical lines
    tblPr.append(borders)

    # Thin rule under the header row (first row)
    if table.rows:
        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcB = tcPr.find(qn("w:tcBorders"))
            if tcB is None:
                tcB = OxmlElement("w:tcBorders")
                tcPr.append(tcB)
            for old in tcB.findall(qn("w:bottom")):
                tcB.remove(old)
            tcB.append(_border("bottom", 8))      # ~1pt thin mid rule

    table.alignment = WD_TABLE_ALIGNMENT.CENTER


# Body-text paragraph styles that should be justified + 1.5-spaced
# (matches the LaTeX \onehalfspacing + justified body + \parskip 0.5em).
BODY_STYLES = {
    "First Paragraph", "Body Text", "Definition", "Definition Term",
    "Compositor", "Bibliography",
}


def center_title_block(doc: Document) -> int:
    """Centre the title block (everything before the abstract). Returns the
    paragraph index where the body begins."""
    end = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        low = p.text.strip().lower()
        if low.startswith("structured abstract") or low == "abstract":
            end = i
            break
    for p in doc.paragraphs[:end]:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return end


def normalize_body(doc: Document, body_start: int) -> int:
    """Justify body paragraphs, set 1.5 line spacing and a uniform 6pt gap;
    give references a hanging indent. Title block (< body_start) is skipped."""
    from docx.shared import Cm
    n = 0
    for i, p in enumerate(doc.paragraphs):
        if i < body_start:
            continue
        if p.style.name not in BODY_STYLES:
            continue  # leave headings, captions, figures untouched
        pf = p.paragraph_format
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        if p.style.name == "Bibliography":
            pf.left_indent = Cm(0.6)          # hanging indent for the
            pf.first_line_indent = Cm(-0.6)   # reference list
        else:
            pf.first_line_indent = Pt(0)
        n += 1
    return n


def set_times_new_roman(path: str) -> None:
    """Force the whole document to Times New Roman (Emerald house font).

    The pandoc reference doc uses the Office theme fonts (Aptos), referenced
    everywhere via minorHAnsi/majorHAnsi. Rewriting the theme's Latin
    typefaces propagates Times New Roman to body, headings and tables in one
    shot. Math (OMML / Cambria Math) is left untouched, as is standard.
    """
    import zipfile, shutil, re as _re
    TNR = 'typeface="Times New Roman" panose="02020603050405020304"'
    tmp = path + ".tmp"
    with zipfile.ZipFile(path, "r") as zin:
        names = zin.namelist()
        data = {n: zin.read(n) for n in names}

    # Patch the theme: replace the major (Aptos Display) and minor (Aptos)
    # Latin typefaces with Times New Roman.
    theme_key = "word/theme/theme1.xml"
    if theme_key in data:
        xml = data[theme_key].decode("utf-8")
        xml = _re.sub(r'typeface="Aptos Display"(\s+panose="[^"]*")?', TNR, xml)
        xml = _re.sub(r'typeface="Aptos"(\s+panose="[^"]*")?', TNR, xml)
        data[theme_key] = xml.encode("utf-8")

    # Belt-and-suspenders: also pin docDefaults to an explicit TNR so any
    # run that does not resolve through the theme still gets Times New Roman.
    styles_key = "word/styles.xml"
    if styles_key in data:
        sx = data[styles_key].decode("utf-8")
        sx = sx.replace(
            '<w:rFonts w:asciiTheme="minorHAnsi" w:cstheme="minorBidi" '
            'w:eastAsiaTheme="minorHAnsi" w:hAnsiTheme="minorHAnsi"/>',
            '<w:rFonts w:ascii="Times New Roman" w:cs="Times New Roman" '
            'w:eastAsia="Times New Roman" w:hAnsi="Times New Roman"/>',
        )
        data[styles_key] = sx.encode("utf-8")

    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, data[n])
    shutil.move(tmp, path)


def main(path: str) -> None:
    doc = Document(path)
    fig_n = tbl_n = 0
    fig_done = tbl_done = 0

    for block in doc.iter_inner_content():
        if isinstance(block, Paragraph):
            style = block.style.name
            if style == "Image Caption":
                fig_n += 1
                prepend_bold_label(block, f"Figure {fig_n}. ")
                block.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fig_done += 1
            elif style == "Table Caption":
                tbl_n += 1
                prepend_bold_label(block, f"Table {tbl_n}. ")
                block.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif style == "Captioned Figure":
                block.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif isinstance(block, Table):
            apply_booktabs(block)
            tbl_done += 1

    # Centre the title block, then justify + 1.5-space the body.
    body_start = center_title_block(doc)
    body_n = normalize_body(doc, body_start)

    # Prepend the editorial note (single source of truth = main.tex)
    note = extract_submitnote()
    if note:
        prepend_note(doc, note)
        print(f"  Editorial note:   '{note[:50]}...'")
    else:
        print("  Editorial note:   (none found in main.tex)")

    doc.save(path)

    # Apply Emerald house font (Times New Roman) by patching the saved file.
    set_times_new_roman(path)

    print(f"  Figures numbered: {fig_n}")
    print(f"  Tables numbered:  {tbl_n}")
    print(f"  Tables restyled:  {tbl_done}")
    print(f"  Title block paras centred: {body_start}")
    print(f"  Body paras justified+1.5:  {body_n}")
    print(f"  Font set to Times New Roman (Emerald house style)")


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else "main.docx")
